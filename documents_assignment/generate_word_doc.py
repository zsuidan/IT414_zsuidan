import PyPDF4
import docx
import openpyxl
from bs4 import BeautifulSoup

#Creates a new document for the report
ceo_report = docx.Document()

#Adds the contents of the operations document to the ceo report
operations_file = docx.Document('text_files/operations.docx')

ceo_report.add_heading("Operations", 1)

    #Adds each paragraph from operations document to the ceo report
for paragraph in operations_file.paragraphs:
    ceo_report.add_paragraph(paragraph.text)

#Adds the contents of the sales spreadsheet to the ceo report
sales_file = openpyxl.load_workbook('text_files/sales.xlsx')

ceo_report.add_heading("Sales", 1)

curr_sheet = sales_file.active
max_row = curr_sheet.max_row
max_col = curr_sheet.max_column

    #Loops through the sales spreadsheet adding rows to the new table
sales_table = ceo_report.add_table(rows=max_row, cols=max_col)

row_count = 1

while row_count <= max_row:
    tmp_row_cells = sales_table.rows[row_count - 1].cells
    col_count = 1

    while col_count <= max_col:
        #Formatting for margins
        if col_count == 5 and row_count > 1:
            tmp_row_cells[col_count - 1].text = str(curr_sheet.cell(row=row_count, column=col_count).value * 100) + "%"
        #Formatting for other numbers
        elif col_count > 1 and row_count > 1:
            #If the number has no commas in it already, they are added, otherwise the number is added normally
            try:
                raw_number = curr_sheet.cell(row=row_count, column=col_count).value
                formatted_number = "{:,}".format(raw_number)
                tmp_row_cells[col_count - 1].text = formatted_number
            except:
                tmp_row_cells[col_count - 1].text = str(curr_sheet.cell(row=row_count, column=col_count).value)
        else:
            tmp_row_cells[col_count - 1].text = str(curr_sheet.cell(row=row_count, column=col_count).value)
        col_count += 1

    row_count += 1

ceo_report.add_page_break()

#Adds the contents of the marketing PDF to the ceo report
marketing_file = open('text_files/marketing.pdf', 'rb')

ceo_report.add_heading("Marketing", 1)

pdfReader = PyPDF4.PdfFileReader(marketing_file)

pdf_page = pdfReader.getPage(0)

    #Pulls text from page and fixes the extraneous spaces and new lines 
pdf_text = pdf_page.extractText().split("\n")

ceo_report.add_paragraph(pdf_text)

marketing_file.close()


#Adds the contents of the IT webpage to the ceo report
it_file = open('text_files/IT.html')
it_soup = BeautifulSoup(it_file, "lxml")

ceo_report.add_heading("IT", 1)

    #Pulls all paragraphs from the webpage then loops through them adding each individually
it_paragraphs = it_soup.findAll("p")

for para in it_paragraphs:
    ceo_report.add_paragraph(para.getText())

#Saves the CEO report document to the text_files folder
ceo_report.save("text_files/ceo_report.docx")